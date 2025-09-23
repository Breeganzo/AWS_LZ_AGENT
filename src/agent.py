import os
import json
import logging
from datetime import datetime
from typing import Dict, List, Any, Optional
from io import BytesIO

# Import modules
from questionnaire import AWSQuestionnaireEngine
from decision_matrix import AWSDecisionMatrix
from diagram_generator import AWSDiagramGenerator

# Try to import optional dependencies
try:
    import google.generativeai as genai
    from dotenv import load_dotenv
    GEMINI_AVAILABLE = True
    load_dotenv()
except ImportError:
    GEMINI_AVAILABLE = False
    genai = None

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    Document = None

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class AWSLandingZoneAgent:
    """
    AWS Landing Zone Architecture Consultant Agent
    
    An intelligent agent that provides AWS architecture recommendations
    based on industry requirements, compliance needs, and business constraints.
    """
    
    def __init__(self):
        """Initialize the AWS Landing Zone Agent"""
        self.questionnaire_engine = AWSQuestionnaireEngine()
        self.decision_matrix = AWSDecisionMatrix()
        self.diagram_generator = AWSDiagramGenerator()
        self.answers = {}
        self.timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        # Initialize Gemini AI if available
        self.gemini_model = None
        if GEMINI_AVAILABLE and genai:
            api_key = os.getenv('GEMINI_API_KEY')
            if api_key:
                try:
                    genai.configure(api_key=api_key)
                    self.gemini_model = genai.GenerativeModel('gemini-1.5-flash')
                    logger.info("Gemini AI initialized successfully")
                except Exception as e:
                    logger.warning(f"Failed to initialize Gemini AI: {e}")
    
    def get_questions(self) -> List[Dict[str, Any]]:
        """Get the list of assessment questions"""
        return self.questionnaire_engine.get_questions()
    
    def save_answer(self, question_id: str, answer: Any):
        """Save an answer to the assessment"""
        self.answers[question_id] = answer
        logger.info(f"Saved answer for question {question_id}: {answer}")
    
    def get_recommendation(self) -> Dict[str, Any]:
        """Generate architecture recommendation based on answers"""
        logger.info("Generating architecture recommendation...")
        
        # Use decision matrix to get recommendations
        recommendation = self.decision_matrix.get_recommendation(self.answers)
        
        # Enhance with AI insights if available
        if self.gemini_model:
            ai_insights = self._get_ai_insights(recommendation, self.answers)
            recommendation['ai_insights'] = ai_insights
        
        return recommendation
    
    def _get_ai_insights(self, recommendation: Dict[str, Any], answers: Dict[str, Any]) -> str:
        """Get AI-powered insights for the recommendation"""
        try:
            prompt = f"""
            As an AWS Solutions Architect, provide expert insights for the following AWS Landing Zone architecture recommendation:

            **Recommended Pattern:** {recommendation['pattern']['name']}
            **Industry:** {answers.get('industry', 'General')}
            **Compliance Requirements:** {answers.get('compliance_requirements', [])}
            **Applications Count:** {answers.get('application_count', 'Unknown')}
            **Security Level:** {answers.get('security_level', 'Standard')}
            **Regions:** {answers.get('regions', [])}

            Please provide:
            1. **Implementation Best Practices** - Key considerations for this specific setup
            2. **Potential Challenges** - Common pitfalls and how to avoid them
            3. **Migration Strategy** - High-level approach for implementation
            4. **Cost Optimization** - Ways to optimize costs for this architecture
            5. **Security Considerations** - Additional security measures to implement

            Keep the response practical, specific to AWS services, and professional.
            """
            
            response = self.gemini_model.generate_content(prompt)
            return response.text
            
        except Exception as e:
            logger.error(f"Error getting AI insights: {e}")
            return "AI insights temporarily unavailable. Please refer to AWS documentation for best practices."
    
    def generate_mermaid_diagram(self, recommendation: Dict[str, Any]) -> str:
        """Generate Mermaid.js diagram"""
        return self.diagram_generator.generate_mermaid(recommendation, self.answers)
    
    def generate_drawio_diagram(self, recommendation: Dict[str, Any]) -> str:
        """Generate Draw.io XML diagram"""
        return self.diagram_generator.generate_drawio(recommendation, self.answers)
    
    def save_diagrams(self, recommendation: Dict[str, Any], diagram_format: str = "both") -> Dict[str, str]:
        """
        Save diagrams to organized directories
        
        Args:
            recommendation: The recommendation dict
            diagram_format: "mermaid", "drawio", or "both"
        
        Returns:
            Dict with paths of saved files
        """
        return self.diagram_generator.save_diagrams(recommendation, diagram_format)
    
    def export_to_json(self, recommendation: Dict[str, Any]) -> str:
        """Export recommendation to JSON format"""
        export_data = {
            "timestamp": self.timestamp,
            "assessment_answers": self.answers,
            "recommendation": recommendation,
            "metadata": {
                "agent_version": "2.0",
                "export_format": "json"
            }
        }
        return json.dumps(export_data, indent=2, default=str)
    
    def export_to_markdown(self, recommendation: Dict[str, Any]) -> str:
        """Export recommendation to Markdown format"""
        md_content = f"""# AWS Landing Zone Architecture Report

**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

## 📊 Assessment Summary

### Recommended Architecture Pattern
**{recommendation['pattern']['name']}**

{recommendation['pattern']['description']}

**Confidence Score:** {recommendation['confidence']:.1%}

## 📋 Your Requirements

"""
        
        # Add assessment answers
        for question_id, answer in self.answers.items():
            question = next((q for q in self.get_questions() if q['id'] == question_id), None)
            if question:
                md_content += f"**{question['question']}**  \n{answer}\n\n"
        
        # Add AI insights if available
        if recommendation.get('ai_insights'):
            md_content += f"""## 🧠 AI-Powered Insights

{recommendation['ai_insights']}

"""
        
        # Add implementation details
        md_content += f"""## 🏗️ Architecture Details

### AWS Services Included
- AWS Control Tower
- AWS Organizations
- AWS IAM Identity Center
- Amazon VPC
- AWS CloudTrail
- AWS Config
- AWS Security Hub

### Organizational Units Structure
- **Security OU**: Centralized security and compliance
- **Production OU**: Live production workloads
- **Non-Production OU**: Development, testing, and staging environments

## 📑 Export Information

- **Report ID:** {self.timestamp}
- **Agent Version:** AWS Landing Zone Architect v2.0
- **Export Format:** Markdown

---
*This report was generated by the AWS Landing Zone Architecture Consultant Agent*
"""
        return md_content
    
    def export_to_word(self, recommendation: Dict[str, Any]) -> bytes:
        """Export recommendation to Word document format"""
        if not DOCX_AVAILABLE or Document is None:
            raise ImportError("python-docx is required for Word export. Please install it with: pip install python-docx")
        
        doc = Document()
        
        # Title
        title = doc.add_heading('AWS Landing Zone Architecture Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtitle with timestamp
        subtitle = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Executive Summary
        doc.add_heading('📊 Executive Summary', level=1)
        
        summary_para = doc.add_paragraph()
        summary_para.add_run('Recommended Architecture: ').bold = True
        summary_para.add_run(recommendation['pattern']['name'])
        
        doc.add_paragraph(recommendation['pattern']['description'])
        
        confidence_para = doc.add_paragraph()
        confidence_para.add_run('Confidence Score: ').bold = True
        confidence_para.add_run(f"{recommendation['confidence']:.1%}")
        
        # Requirements section
        doc.add_heading('📋 Assessment Responses', level=1)
        
        for question_id, answer in self.answers.items():
            question = next((q for q in self.get_questions() if q['id'] == question_id), None)
            if question:
                doc.add_paragraph(question['question'], style='Heading 3')
                doc.add_paragraph(str(answer))
        
        # AI Insights
        if recommendation.get('ai_insights'):
            doc.add_heading('🧠 AI-Powered Implementation Insights', level=1)
            doc.add_paragraph(recommendation['ai_insights'])
        
        # Architecture Details
        doc.add_heading('🏗️ Architecture Implementation Details', level=1)
        
        doc.add_paragraph('Core AWS Services:', style='Heading 3')
        services = [
            'AWS Control Tower - Landing zone setup and governance',
            'AWS Organizations - Multi-account management',
            'AWS IAM Identity Center - Centralized identity management',
            'Amazon VPC - Network isolation and security',
            'AWS CloudTrail - API logging and audit trails',
            'AWS Config - Configuration compliance monitoring',
            'AWS Security Hub - Centralized security findings'
        ]
        
        for service in services:
            doc.add_paragraph(service, style='List Bullet')
        
        # Footer
        doc.add_page_break()
        footer_para = doc.add_paragraph()
        footer_para.add_run('Generated by: ').italic = True
        footer_para.add_run('AWS Landing Zone Architecture Consultant Agent v2.0').italic = True
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save to BytesIO
        doc_buffer = BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        
        return doc_buffer.read()
    
    def get_timestamp(self) -> str:
        """Get the current timestamp for file naming"""
        return self.timestamp